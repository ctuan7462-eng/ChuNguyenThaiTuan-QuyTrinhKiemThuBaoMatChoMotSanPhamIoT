#!/usr/bin/env python3
"""
=============================================================
IoT Security Testing Report Generator
Đề tài 50 – Quy trình kiểm thử bảo mật cho một sản phẩm IoT
Sinh viên: Chu Nguyễn Thái Tuấn – MSSV: 231A390004
=============================================================

Cách dùng:
  1. Chỉnh dữ liệu checklist và findings bên dưới (hoặc load từ file JSON)
  2. Chạy: python generate_report.py
  3. File báo cáo sẽ được xuất ra: BaoCao_KiemThu_IoT.docx

Cấu trúc dữ liệu:
  - PROJECT_INFO   : thông tin đề tài / sản phẩm
  - CHECKLIST      : 25 tiêu chí ánh xạ OWASP ISVS
  - FINDINGS       : danh sách lỗ hổng phát hiện
"""

import json
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# 1. DỮ LIỆU ĐẦU VÀO – chỉnh ở đây hoặc load từ JSON
# ─────────────────────────────────────────────────────────────

PROJECT_INFO = {
    "product_name": "Camera IP Gia Đình (Giả Định)",
    "product_version": "FW-v2.3.1-sim",
    "tester": "Chu Nguyễn Thái Tuấn",
    "mssv": "231A390004",
    "lop": "253INT441001",
    "gv_huong_dan": "ThS. Hồ Nhựt Minh",
    "repo": "https://github.com/ctuan7462-eng/ChuNguyenThaiTuan-QuyTrinhKiemThuChoMotSanPhamIoT",
    "test_date": datetime.now().strftime("%d/%m/%Y"),
    "report_date": datetime.now().strftime("%d/%m/%Y"),
    "ghi_chu": "Toàn bộ dữ liệu là giả lập – mục đích minh họa quy trình, không phải kết quả kiểm thử thực tế.",
}

CHECKLIST = [
    # format: (id, nhom, test_point, isvs_ref, ket_qua)
    # ket_qua: "Đạt" | "Chưa đạt" | "N/A"
    ("TC-01", "Phần cứng", "Cổng UART/JTAG có bị khóa ở bản production", "Hardware security", "Chưa đạt"),
    ("TC-02", "Phần cứng", "Thiết bị có cơ chế tamper-evidence (niêm phong)", "Hardware security", "N/A"),
    ("TC-03", "Phần cứng", "Không có nhãn/label lộ thông tin nhạy cảm", "Hardware security", "Đạt"),
    ("TC-04", "Phần cứng", "microSD/khe cắm ngoài có được bảo vệ", "Hardware security", "Chưa đạt"),
    ("TC-05", "Firmware", "Không có credential hard-code trong firmware", "Secure credential storage", "Chưa đạt"),
    ("TC-06", "Firmware", "Firmware được ký số trước khi phân phối", "Firmware integrity", "Chưa đạt"),
    ("TC-07", "Firmware", "Dịch vụ Telnet/SSH mặc định bị tắt", "Secure configuration", "Chưa đạt"),
    ("TC-08", "Firmware", "Log hệ thống không ghi dữ liệu nhạy cảm dạng rõ", "Data protection", "Đạt"),
    ("TC-09", "Mạng",     "Giao tiếp camera–cloud dùng TLS 1.2+", "Encryption in transit", "Đạt"),
    ("TC-10", "Mạng",     "Server có được xác thực (cert validation) ở client", "Mutual authentication", "Chưa đạt"),
    ("TC-11", "Mạng",     "RTSP/stream video có mã hóa", "Encryption in transit", "Chưa đạt"),
    ("TC-12", "Mạng",     "Có cơ chế chống replay cho lệnh điều khiển", "Anti-replay", "Chưa đạt"),
    ("TC-13", "Mạng",     "Không có cổng quản trị mở ra Internet trực tiếp", "Network exposure", "Chưa đạt"),
    ("TC-14", "API/Cloud","API xác thực bằng token/JWT có thời hạn", "Authentication", "Đạt"),
    ("TC-15", "API/Cloud","API kiểm tra quyền theo device_id (BOLA check)", "Authorization", "Chưa đạt"),
    ("TC-16", "API/Cloud","Có giới hạn tần suất request (rate limiting)", "Anti-automation", "Chưa đạt"),
    ("TC-17", "API/Cloud","Response API không trả thừa thông tin nội bộ", "Data protection", "Đạt"),
    ("TC-18", "Mobile App","Token lưu trong Keystore/Keychain, không plaintext", "Secure storage", "Đạt"),
    ("TC-19", "Mobile App","App có kiểm tra certificate hợp lệ (không bypass TLS)", "Communication security", "Chưa đạt"),
    ("TC-20", "Mobile App","Quyền app yêu cầu tối thiểu cần thiết", "Least privilege", "Chưa đạt"),
    ("TC-21", "Update",   "Firmware update được kiểm tra hash/chữ ký trước khi cài", "Secure update", "Chưa đạt"),
    ("TC-22", "Update",   "Có cơ chế rollback về bản trước nếu update lỗi", "Update integrity", "N/A"),
    ("TC-23", "Update",   "Kênh tải update có mã hóa (HTTPS)", "Encryption in transit", "Đạt"),
    ("TC-24", "Update",   "Thiết bị kiểm tra phiên bản, không hạ cấp trái phép", "Anti-rollback attack", "Chưa đạt"),
    ("TC-25", "Vận hành", "Có tài liệu hướng dẫn đổi mật khẩu mặc định lần đầu", "Default credentials", "Đạt"),
]

FINDINGS = [
    {
        "id": "FIND-001",
        "pham_vi": "Firmware",
        "mo_ta": (
            "Phát hiện chuỗi credential dạng rõ trong file cấu hình "
            "firmware giả lập (/etc/config/system.conf). Mật khẩu admin "
            "mặc định giống nhau trên toàn dòng sản phẩm."
        ),
        "bang_chung": "Chuỗi hard-code: admin:password123 trong file cấu hình giả lập",
        "muc_do": "Cao",
        "kha_nang": "Trung bình",
        "khuyen_nghi": (
            "Loại bỏ toàn bộ hard-code credential. Áp dụng cấp key "
            "riêng từng thiết bị (unique per-device key) khi xuất xưởng. "
            "Bắt người dùng đổi mật khẩu lần đầu đăng nhập."
        ),
        "trang_thai": "Mở",
        "tc_lien_quan": "TC-05",
    },
    {
        "id": "FIND-002",
        "pham_vi": "Mạng",
        "mo_ta": (
            "Luồng RTSP video truyền không mã hóa trong mạng nội bộ "
            "(rtsp://192.168.x.x:554/stream). Kẻ tấn công trong cùng "
            "mạng Wi-Fi có thể capture và xem trực tiếp."
        ),
        "bang_chung": "Gói RTSP bắt được bằng Wireshark (mô phỏng) – payload video rõ không mã hóa",
        "muc_do": "Trung bình",
        "kha_nang": "Cao",
        "khuyen_nghi": (
            "Áp dụng SRTP hoặc tunnel luồng video qua TLS. "
            "Xác thực người xem trước khi cấp quyền truy cập stream."
        ),
        "trang_thai": "Mở",
        "tc_lien_quan": "TC-11",
    },
    {
        "id": "FIND-003",
        "pham_vi": "API/Cloud",
        "mo_ta": (
            "Endpoint GET /api/v1/device/{device_id}/stream không kiểm "
            "tra quyền sở hữu thiết bị. Người dùng A có thể truy cập "
            "stream của thiết bị thuộc tài khoản B bằng cách thay device_id."
        ),
        "bang_chung": "Yêu cầu API giả lập: GET /api/v1/device/CAM-99999/stream trả về 200 OK với token của người dùng khác",
        "muc_do": "Nghiêm trọng",
        "kha_nang": "Trung bình",
        "khuyen_nghi": (
            "Bổ sung kiểm tra ownership ở tầng authorization: mỗi request "
            "phải xác minh device_id thuộc về user_id trong token. "
            "Áp dụng policy ABAC (Attribute-Based Access Control)."
        ),
        "trang_thai": "Mở",
        "tc_lien_quan": "TC-15",
    },
    {
        "id": "FIND-004",
        "pham_vi": "Mobile App",
        "mo_ta": (
            "App bỏ qua lỗi certificate không hợp lệ – chấp nhận "
            "self-signed certificate bất kỳ. Cho phép thực hiện MITM "
            "attack để đọc/sửa dữ liệu giữa app và cloud server."
        ),
        "bang_chung": "Thử nghiệm giả lập với proxy MITM – app kết nối thành công khi dùng cert tự ký",
        "muc_do": "Cao",
        "kha_nang": "Thấp",
        "khuyen_nghi": (
            "Bật certificate pinning (hoặc public key pinning) trong app. "
            "Từ chối kết nối khi cert không khớp với fingerprint đã pin."
        ),
        "trang_thai": "Mở",
        "tc_lien_quan": "TC-19",
    },
    {
        "id": "FIND-005",
        "pham_vi": "Update",
        "mo_ta": (
            "Cơ chế OTA tải và cài firmware mà không kiểm tra chữ ký "
            "số hoặc hash. Kẻ tấn công kiểm soát kênh phân phối update "
            "có thể push firmware độc hại xuống toàn bộ thiết bị."
        ),
        "bang_chung": "Quy trình OTA giả lập chấp nhận file firmware không có chữ ký – cài đặt thành công",
        "muc_do": "Nghiêm trọng",
        "kha_nang": "Thấp",
        "khuyen_nghi": (
            "Ký số firmware bằng private key của hãng (RSA-2048 hoặc ECDSA). "
            "Thiết bị xác minh chữ ký trước khi flash. "
            "Bổ sung cơ chế rollback an toàn về bản firmware đã xác minh."
        ),
        "trang_thai": "Mở",
        "tc_lien_quan": "TC-21",
    },
]

# ─────────────────────────────────────────────────────────────
# 2. LOAD DỮ LIỆU TỪ FILE JSON (tuỳ chọn)
# ─────────────────────────────────────────────────────────────

def load_from_json(json_path: str):
    """Load PROJECT_INFO, CHECKLIST, FINDINGS từ file JSON ngoài."""
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    info = data.get("project_info", PROJECT_INFO)
    cl = [tuple(item) for item in data.get("checklist", [])]
    fi = data.get("findings", [])
    return info, cl, fi


# ─────────────────────────────────────────────────────────────
# 3. HELPER: THỐNG KÊ
# ─────────────────────────────────────────────────────────────

def stats(checklist):
    total = len(checklist)
    dat   = sum(1 for r in checklist if r[4] == "Đạt")
    chua  = sum(1 for r in checklist if r[4] == "Chưa đạt")
    na    = sum(1 for r in checklist if r[4] == "N/A")
    return total, dat, chua, na


def risk_score(muc_do: str, kha_nang: str) -> int:
    """Tính điểm rủi ro để sắp xếp finding theo mức ưu tiên."""
    sev = {"Nghiêm trọng": 4, "Cao": 3, "Trung bình": 2, "Thấp": 1}
    prob = {"Cao": 3, "Trung bình": 2, "Thấp": 1}
    return sev.get(muc_do, 0) * prob.get(kha_nang, 0)


# ─────────────────────────────────────────────────────────────
# 4. HELPER: ĐỊNH DẠNG DOCX
# ─────────────────────────────────────────────────────────────

FONT_NAME = "Times New Roman"
COLOR_HEADER = RGBColor(0x1F, 0x49, 0x7D)   # xanh đậm
COLOR_PASS   = RGBColor(0x37, 0x86, 0x44)   # xanh lá
COLOR_FAIL   = RGBColor(0xC0, 0x00, 0x00)   # đỏ
COLOR_NA     = RGBColor(0x80, 0x80, 0x80)   # xám
COLOR_CRIT   = RGBColor(0xC0, 0x00, 0x00)
COLOR_HIGH   = RGBColor(0xFF, 0x00, 0x00)
COLOR_MED    = RGBColor(0xFF, 0x92, 0x00)
COLOR_LOW    = RGBColor(0x00, 0x70, 0xC0)

SEVERITY_COLOR = {
    "Nghiêm trọng": COLOR_CRIT,
    "Cao": COLOR_HIGH,
    "Trung bình": COLOR_MED,
    "Thấp": COLOR_LOW,
}

STATUS_COLOR = {
    "Đạt":       COLOR_PASS,
    "Chưa đạt":  COLOR_FAIL,
    "N/A":       COLOR_NA,
}


def set_cell_bg(cell, hex_color: str):
    """Tô màu nền cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
                  color=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    p = add_paragraph(doc, text, bold=True, size=sizes.get(level, 12),
                      color=COLOR_HEADER, space_after=4)
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        # gạch dưới heading 1
        pPr = p.runs[0]._element.get_or_add_rPr()
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        pPr.append(u)
    return p


def make_table(doc, headers, rows, col_widths_cm):
    """Tạo bảng với header màu xanh đậm, font Times New Roman."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Table Grid"

    # header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # data rows
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, (text, extra) in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = extra.get("align", WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(str(text))
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            if "color" in extra:
                run.font.color.rgb = extra["color"]
            if extra.get("bold"):
                run.bold = True
    return tbl


# ─────────────────────────────────────────────────────────────
# 5. SINH BÁO CÁO
# ─────────────────────────────────────────────────────────────

def generate_report(info, checklist, findings, out_path="BaoCao_KiemThu_IoT.docx"):
    doc = Document()

    # --- margin ---
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # ── TRANG BÌA ────────────────────────────────────────────
    add_paragraph(doc, "TRƯỜNG ĐẠI HỌC VĂN HIẾN", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "KHOA CÔNG NGHỆ THÔNG TIN", bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_paragraph(doc, "BÁO CÁO KIỂM THỬ BẢO MẬT IoT", bold=True, size=18,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=COLOR_HEADER, space_after=4)
    add_paragraph(doc, "Đề tài 50: Quy trình kiểm thử bảo mật cho một sản phẩm IoT",
                  bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # bảng thông tin
    info_rows = [
        [("Sản phẩm kiểm thử", {}), (info["product_name"], {"bold": True})],
        [("Phiên bản firmware",  {}), (info["product_version"], {})],
        [("Người thực hiện",     {}), (f'{info["tester"]}  –  MSSV: {info["mssv"]}', {})],
        [("Lớp học phần",        {}), (info["lop"], {})],
        [("Giảng viên hướng dẫn",{}), (info["gv_huong_dan"], {})],
        [("Ngày kiểm thử",       {}), (info["test_date"], {})],
        [("Ngày xuất báo cáo",   {}), (info["report_date"], {})],
        [("Repo GitHub",         {}), (info["repo"], {"color": RGBColor(0, 0, 0xFF)})],
    ]
    make_table(doc, ["Trường", "Nội dung"], info_rows, [5, 12])

    doc.add_paragraph()
    add_paragraph(doc, f'⚠  {info["ghi_chu"]}', italic=True, size=10,
                  color=COLOR_NA, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ── 1. THỐNG KÊ TỔNG QUAN ────────────────────────────────
    add_heading(doc, "1. TỔNG QUAN KẾT QUẢ KIỂM THỬ", level=1)
    total, dat, chua, na = stats(checklist)
    ty_le = round(dat / (total - na) * 100, 1) if (total - na) > 0 else 0
    so_finding = len(findings)
    nghiem_trong = sum(1 for f in findings if f["muc_do"] == "Nghiêm trọng")
    cao          = sum(1 for f in findings if f["muc_do"] == "Cao")

    summary_rows = [
        [("Tổng số tiêu chí kiểm thử",    {}), (str(total),             {"align": WD_ALIGN_PARAGRAPH.CENTER})],
        [("Đạt",                           {}), (str(dat),               {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": COLOR_PASS, "bold": True})],
        [("Chưa đạt",                      {}), (str(chua),              {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": COLOR_FAIL, "bold": True})],
        [("Không áp dụng (N/A)",           {}), (str(na),                {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": COLOR_NA})],
        [("Tỷ lệ đạt (loại trừ N/A)",     {}), (f"{ty_le}%",            {"align": WD_ALIGN_PARAGRAPH.CENTER, "bold": True})],
        [("Tổng số lỗ hổng phát hiện",    {}), (str(so_finding),        {"align": WD_ALIGN_PARAGRAPH.CENTER, "bold": True})],
        [("  — Mức Nghiêm trọng",          {}), (str(nghiem_trong),      {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": COLOR_CRIT, "bold": True})],
        [("  — Mức Cao",                   {}), (str(cao),               {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": COLOR_HIGH, "bold": True})],
    ]
    make_table(doc, ["Chỉ số", "Giá trị"], summary_rows, [10, 7])
    doc.add_paragraph()

    # ── 2. CHECKLIST ─────────────────────────────────────────
    add_heading(doc, "2. CHECKLIST YÊU CẦU KIỂM THỬ (25 TIÊU CHÍ – OWASP ISVS)", level=1)
    add_paragraph(doc,
        "Mỗi tiêu chí được ánh xạ với yêu cầu tương ứng trong OWASP IoT Security Verification Standard (ISVS). "
        "Kết quả: Đạt / Chưa đạt / N/A.", size=10, italic=True)

    cl_rows = []
    prev_group = None
    for (tid, nhom, test_point, isvs, ket_qua) in checklist:
        kq_color = STATUS_COLOR.get(ket_qua, COLOR_NA)
        group_cell = (nhom, {"bold": True}) if nhom != prev_group else ("", {})
        prev_group = nhom
        cl_rows.append([
            (tid,        {"align": WD_ALIGN_PARAGRAPH.CENTER}),
            group_cell,
            (test_point, {}),
            (isvs,       {"color": COLOR_HEADER}),
            (ket_qua,    {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": kq_color, "bold": True}),
        ])

    make_table(doc,
               ["ID", "Nhóm", "Tiêu chí kiểm thử", "OWASP ISVS", "Kết quả"],
               cl_rows,
               [1.4, 2.4, 7.5, 4.0, 2.0])
    doc.add_paragraph()

    # ── 3. FINDINGS ──────────────────────────────────────────
    add_heading(doc, "3. BÁO CÁO PHÁT HIỆN LỖ HỔNG (VULNERABILITY FINDINGS)", level=1)

    sorted_findings = sorted(findings,
                             key=lambda f: risk_score(f["muc_do"], f["kha_nang"]),
                             reverse=True)

    for fi in sorted_findings:
        md_color = SEVERITY_COLOR.get(fi["muc_do"], COLOR_NA)

        add_heading(doc,
            f'[{fi["id"]}]  {fi["pham_vi"]}  —  Mức: {fi["muc_do"]}  |  Khả năng: {fi["kha_nang"]}',
            level=2)

        fi_rows = [
            [("Phạm vi / Thành phần",  {}), (fi["pham_vi"], {"bold": True})],
            [("Mô tả phát hiện",       {}), (fi["mo_ta"],   {})],
            [("Bằng chứng (giả lập)",  {}), (fi["bang_chung"], {"color": COLOR_NA, "bold": False})],
            [("Mức độ nghiêm trọng",   {}), (fi["muc_do"],  {"color": md_color, "bold": True,
                                                              "align": WD_ALIGN_PARAGRAPH.CENTER})],
            [("Khả năng xảy ra",       {}), (fi["kha_nang"],{"align": WD_ALIGN_PARAGRAPH.CENTER})],
            [("Tiêu chí liên quan",    {}), (fi["tc_lien_quan"], {"color": COLOR_HEADER})],
            [("Khuyến nghị",           {}), (fi["khuyen_nghi"], {"bold": False})],
            [("Trạng thái",            {}), (fi["trang_thai"], {"bold": True,
                                                                 "align": WD_ALIGN_PARAGRAPH.CENTER,
                                                                 "color": COLOR_FAIL if fi["trang_thai"] == "Mở" else COLOR_PASS})],
        ]
        make_table(doc, ["Trường", "Chi tiết"], fi_rows, [4, 13])
        doc.add_paragraph()

    # ── 4. MA TRẬN RỦI RO ────────────────────────────────────
    add_heading(doc, "4. MA TRẬN MỨC ĐỘ RỦI RO VÀ ƯU TIÊN XỬ LÝ", level=1)

    # bảng ưu tiên
    prio_rows = []
    for rank, fi in enumerate(sorted_findings, 1):
        md_color = SEVERITY_COLOR.get(fi["muc_do"], COLOR_NA)
        prio_rows.append([
            (str(rank),       {"align": WD_ALIGN_PARAGRAPH.CENTER, "bold": True}),
            (fi["id"],        {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": COLOR_HEADER, "bold": True}),
            (fi["pham_vi"],   {"align": WD_ALIGN_PARAGRAPH.CENTER}),
            (fi["muc_do"],    {"align": WD_ALIGN_PARAGRAPH.CENTER, "color": md_color, "bold": True}),
            (fi["kha_nang"],  {"align": WD_ALIGN_PARAGRAPH.CENTER}),
            (str(risk_score(fi["muc_do"], fi["kha_nang"])),
                              {"align": WD_ALIGN_PARAGRAPH.CENTER, "bold": True}),
            (fi["trang_thai"],{"align": WD_ALIGN_PARAGRAPH.CENTER}),
        ])

    make_table(doc,
               ["#", "Finding", "Thành phần", "Mức độ", "Khả năng", "Điểm RR", "Trạng thái"],
               prio_rows,
               [0.8, 2.0, 2.5, 2.5, 2.5, 2.0, 2.7])
    doc.add_paragraph()

    add_paragraph(doc,
        "Điểm rủi ro = Mức độ × Khả năng  "
        "(Nghiêm trọng=4, Cao=3, TB=2, Thấp=1  ×  Cao=3, TB=2, Thấp=1). "
        "Điểm càng cao → ưu tiên xử lý càng sớm.",
        italic=True, size=9, color=COLOR_NA)

    # ── 5. CHÚ THÍCH & TÀI LIỆU THAM KHẢO ──────────────────
    doc.add_page_break()
    add_heading(doc, "5. TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "OWASP Foundation. OWASP IoT Security Testing Guide (ISTG). https://github.com/OWASP/owasp-istg",
        "OWASP Foundation. OWASP IoT Security Verification Standard (ISVS). https://github.com/OWASP/IoT-Security-Verification-Standard-ISVS",
        "OWASP Foundation. OWASP Firmware Security Testing Methodology (FSTM). https://github.com/scriptingxss/owasp-fstm",
        "OWASP Foundation. OWASP IoTGoat Project. https://github.com/OWASP/IoTGoat",
        "Brian Russell, Drew Van Duren. Practical Internet of Things Security. Packt Publishing, 2016.",
        f"Chu Nguyễn Thái Tuấn. Repo GitHub đề tài 50. {info['repo']}",
    ]
    for i, ref in enumerate(refs, 1):
        add_paragraph(doc, f"[{i}] {ref}", size=10, space_after=4)

    add_paragraph(doc)
    add_paragraph(doc,
        f'Báo cáo được sinh tự động bởi generate_report.py  —  {info["report_date"]}',
        size=9, italic=True, color=COLOR_NA, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(out_path)
    return out_path


# ─────────────────────────────────────────────────────────────
# 6. ENTRY POINT
# ─────────────────────────────────────────────────────────────

def main():
    info, checklist, findings = PROJECT_INFO, CHECKLIST, FINDINGS

    # Nếu truyền file JSON qua dòng lệnh: python generate_report.py data.json
    if len(sys.argv) > 1 and os.path.isfile(sys.argv[1]):
        print(f"[INFO] Load dữ liệu từ: {sys.argv[1]}")
        info, checklist, findings = load_from_json(sys.argv[1])

    out = "BaoCao_KiemThu_IoT.docx"
    print("[INFO] Đang sinh báo cáo...")
    generate_report(info, checklist, findings, out_path=out)

    total, dat, chua, na = stats(checklist)
    print(f"[OK]  Báo cáo đã xuất: {out}")
    print(f"      Checklist: {total} tiêu chí  |  Đạt: {dat}  |  Chưa đạt: {chua}  |  N/A: {na}")
    print(f"      Findings:  {len(findings)}  |  "
          f"Nghiêm trọng: {sum(1 for f in findings if f['muc_do']=='Nghiêm trọng')}  |  "
          f"Cao: {sum(1 for f in findings if f['muc_do']=='Cao')}")


if __name__ == "__main__":
    main()
